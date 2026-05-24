from linebot.v3 import WebhookHandler
from linebot.v3.messaging import (
    Configuration, ApiClient, MessagingApi, MessagingApiBlob,
    ReplyMessageRequest, PushMessageRequest, TextMessage, AudioMessage,
    QuickReply, QuickReplyItem, MessageAction,
    FlexMessage, FlexContainer
)
from linebot.v3.webhooks import MessageEvent, TextMessageContent, AudioMessageContent
from linebot.v3.exceptions import InvalidSignatureError

from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnableParallel, RunnablePassthrough, RunnableLambda
from langchain_community.document_loaders import DirectoryLoader, TextLoader, PyPDFLoader
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_ollama import ChatOllama
from langchain_text_splitters import RecursiveCharacterTextSplitter

from fastapi import FastAPI, Request, HTTPException
from fastapi.responses import JSONResponse
from fastapi.staticfiles import StaticFiles

from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload

from opencc import OpenCC
import edge_tts
from pydub import AudioSegment
from pydub.utils import mediainfo

import os
import re
import asyncio
from dotenv import load_dotenv

load_dotenv()
import uvicorn
import whisper
import io
import shutil
import json
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path
from datetime import datetime

os.environ["TOKENIZERS_PARALLELISM"] = "false"

# ========== 日誌輸出函數 ==========

def log(message, level="INFO"):
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    print(f"[{timestamp}] [{level}] {message}")

# 初始化 FastAPI app
log("正在初始化 FastAPI 應用...", "系統")
app = FastAPI()

# 確保資料夾存在
os.makedirs("audio", exist_ok=True)
os.makedirs("drive_documents", exist_ok=True)
os.makedirs("breeze_vectorstores", exist_ok=True)
log("資料夾建立完成", "系統")

# 註冊靜態資料夾
app.mount("/audio", StaticFiles(directory="audio"), name="audio")

handler = WebhookHandler(channel_secret=os.getenv("LINE_CHANNEL_SECRET"))
configuration = Configuration(access_token=os.getenv("LINE_CHANNEL_ACCESS_TOKEN"))

ngrok_url = os.getenv("NGROK_URL")

# ========== Google Drive / Sheets 設定 ==========

log("正在連接 Google Drive / Sheets API...", "系統")
GOOGLE_DRIVE_FOLDER_ID = os.getenv("GOOGLE_DRIVE_FOLDER_ID")
SHARED_VECTORSTORE_ID  = "shared"
SERVICE_ACCOUNT_FILE   = "service_account.json"

SPREADSHEET_ID       = os.getenv("GOOGLE_SPREADSHEET_ID")
RATINGS_SHEET_NAME   = "breeze_評分記錄"
NO_RESULT_SHEET_NAME = "breeze_查無資料記錄"

SCOPES = [
    'https://www.googleapis.com/auth/drive.readonly',
    'https://www.googleapis.com/auth/spreadsheets',
]

credentials = service_account.Credentials.from_service_account_file(
    SERVICE_ACCOUNT_FILE, scopes=SCOPES)
drive_service  = build('drive',  'v3', credentials=credentials)
sheets_service = build('sheets', 'v4', credentials=credentials)
log("Google Drive / Sheets API 連接成功", "系統")

# ========== Google Sheets 初始化與寫入函數 ==========

def init_sheets():
    try:
        spreadsheet = sheets_service.spreadsheets().get(
            spreadsheetId=SPREADSHEET_ID
        ).execute()
        existing_titles = [s['properties']['title'] for s in spreadsheet['sheets']]

        add_requests = []
        for sheet_name in [RATINGS_SHEET_NAME, NO_RESULT_SHEET_NAME]:
            if sheet_name not in existing_titles:
                add_requests.append({'addSheet': {'properties': {'title': sheet_name}}})

        if add_requests:
            sheets_service.spreadsheets().batchUpdate(
                spreadsheetId=SPREADSHEET_ID,
                body={'requests': add_requests}
            ).execute()
            log(f"已新增工作表: {[r['addSheet']['properties']['title'] for r in add_requests]}", "Sheets")

        headers_map = {
            RATINGS_SHEET_NAME:   [["時間戳記", "使用者ID(前8碼)", "評分"]],
            NO_RESULT_SHEET_NAME: [["時間戳記", "使用者ID(前8碼)", "查無資料的問題", "RAG原始回應", "最終回覆內容（Fallback）"]],
        }
        for sheet_name, headers in headers_map.items():
            result = sheets_service.spreadsheets().values().get(
                spreadsheetId=SPREADSHEET_ID,
                range=f"{sheet_name}!A1"
            ).execute()
            if not result.get('values'):
                sheets_service.spreadsheets().values().update(
                    spreadsheetId=SPREADSHEET_ID,
                    range=f"{sheet_name}!A1",
                    valueInputOption="USER_ENTERED",
                    body={"values": headers}
                ).execute()
                log(f"已寫入標頭: {sheet_name}", "Sheets")

        log("Google Sheets 初始化完成", "系統")
    except Exception as e:
        log(f"Google Sheets 初始化失敗: {e}", "錯誤")


def save_rating_to_sheet(user_id: str, score: str, last_question: str = ""):
    try:
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        row = [[timestamp, user_id[:8] + "...", score]]
        sheets_service.spreadsheets().values().append(
            spreadsheetId=SPREADSHEET_ID,
            range=f"{RATINGS_SHEET_NAME}!A:C",
            valueInputOption="USER_ENTERED",
            body={"values": row}
        ).execute()
        log(f"評分已記錄到 Sheets: {score}", "評分")
    except Exception as e:
        log(f"評分記錄失敗: {e}", "錯誤")


def save_no_result_log(user_id: str, question: str, rag_raw: str = "", fallback_reply: str = ""):
    try:
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        row = [[timestamp, user_id[:8] + "...", question, rag_raw[:200], fallback_reply]]
        sheets_service.spreadsheets().values().append(
            spreadsheetId=SPREADSHEET_ID,
            range=f"{NO_RESULT_SHEET_NAME}!A:E",
            valueInputOption="USER_ENTERED",
            body={"values": row}
        ).execute()
        log(f"查無資料問題已記錄（含Fallback回覆）: {question[:30]}...", "記錄")
    except Exception as e:
        log(f"查無資料記錄失敗: {e}", "錯誤")

# ========== 對話歷史管理 ==========

user_histories: dict[str, list] = {}
user_last_question: dict[str, str] = {}

def get_user_history(user_id: str) -> list:
    return user_histories.get(user_id, [])

def add_to_history(user_id: str, role: str, content: str):
    if user_id not in user_histories:
        user_histories[user_id] = []
    user_histories[user_id].append({"role": role, "content": content})
    if len(user_histories[user_id]) > 10:
        user_histories[user_id] = user_histories[user_id][-10:]

def format_history_text(history: list) -> str:
    lines = []
    for h in history:
        role_label = "使用者" if h["role"] == "user" else "助理"
        lines.append(f"{role_label}：{h['content']}")
    return "\n".join(lines)

# ========== 意圖分類 + 閒聊模板 ==========

CHITCHAT_PATTERNS = [
    (r"謝謝|感謝|多謝|thank",
     "不客氣！有任何問題都歡迎繼續詢問，小伊隨時在這裡 😊"),
    (r"^好的?$|^好啊$|^好喔$|^OK$|^ok$",
     "好的！還有什麼小伊可以幫您的嗎？"),
    (r"還有嗎\??$|還有什麼\??$",
     "您還有其他想了解的嗎？請隨時告訴小伊！😄"),
    (r"^對$|^是的?$|^沒錯$|^嗯$|^嗯嗯$",
     "很高興能幫到您！還有什麼需要了解的嗎？"),
    (r"你好|您好|妳好|哈囉|嗨|^hi$|^hello$",
     "您好！我是小伊，請問有什麼是小伊可以幫您的嗎？😊"),
    (r"^掰掰$|^再見$|^bye$",
     "再見！祝您身體健康，有需要隨時回來找小伊 😊"),
]

FOLLOW_UP_KEYWORDS = ["那", "所以", "那麼", "那這樣", "繼續", "還有", "另外", "剛才", "前面"]

# ★ 方向一：情緒支持關鍵詞（優先於閒聊判斷）
EMOTIONAL_PATTERN = re.compile(
    r"好累|心好累|心累|身心俱疲|撐不住|快撐不下去|受不了|崩潰|喘不過氣|"
    r"好難過|難過|好沮喪|沮喪|好痛苦|很痛苦|"
    r"好絕望|絕望|好無助|無助|好孤單|很孤單|"
    r"心情不好|心情很差|心情低落|情緒低落|"
    r"好恐慌|恐慌發作|好委屈|很委屈|壓垮了|"
    r"不想活|想死|活不下去|好想哭|一直哭",
    re.IGNORECASE
)

def classify_intent(msg: str, history: list) -> str:
    stripped = msg.strip()
    # ★ 情緒支持優先判斷，命中後仍走 RAG
    if EMOTIONAL_PATTERN.search(stripped):
        return "emotional_support"
    for pattern, _ in CHITCHAT_PATTERNS:
        if re.search(pattern, stripped, re.IGNORECASE):
            return "chitchat"
    if history and len(stripped) < 25:
        if any(kw in stripped for kw in FOLLOW_UP_KEYWORDS):
            return "follow_up"
    return "question"

def get_chitchat_response(msg: str) -> str:
    for pattern, reply in CHITCHAT_PATTERNS:
        if re.search(pattern, msg.strip(), re.IGNORECASE):
            return reply
    return "您好！請問有什麼我可以幫您的嗎？😊"

# ★ 方向三：TTS emoji 過濾函數
def strip_emoji(text: str) -> str:
    """移除文字中的 emoji，供 TTS 語音合成使用；TextMessage 仍保留 emoji"""
    emoji_pattern = re.compile("["
        u"\U0001F600-\U0001F64F"
        u"\U0001F300-\U0001F5FF"
        u"\U0001F680-\U0001F6FF"
        u"\U0001F1E0-\U0001F1FF"
        u"\U00002702-\U000027B0"
        u"\U000024C2-\U0001F251"
        u"\U0001f926-\U0001f937"
        u"\U00010000-\U0010ffff"
        u"\u2640-\u2642"
        u"\u2600-\u2B55"
        u"\u200d\u23cf\u23e9\u231a\ufe0f\u3030"
        "]+", flags=re.UNICODE)
    return emoji_pattern.sub('', text).strip()

# ========== Google Drive 文件下載函數（支援分頁，取得全部檔案）==========

def list_all_files_from_gdrive(folder_id: str) -> list:
    """使用 nextPageToken 分頁，確保取回資料夾內所有檔案（突破 100 筆限制）"""
    all_files = []
    page_token = None
    page_num = 0
    query = f"'{folder_id}' in parents and trashed=false"

    while True:
        page_num += 1
        kwargs = dict(
            q=query,
            pageSize=100,
            fields="nextPageToken, files(id, name, mimeType, shortcutDetails)",
        )
        if page_token:
            kwargs["pageToken"] = page_token

        response = drive_service.files().list(**kwargs).execute()
        batch = response.get("files", [])

        resolved_batch = []
        for f in batch:
            if f.get("mimeType") == "application/vnd.google-apps.shortcut":
                target_id = (f.get("shortcutDetails") or {}).get("targetId")
                if target_id:
                    try:
                        real = drive_service.files().get(
                            fileId=target_id,
                            fields="id, name, mimeType"
                        ).execute()
                        log(f"捷徑解析：{f['name']} -> {real['name']} ({real['mimeType']})", "Google Drive")
                        resolved_batch.append(real)
                    except Exception as e:
                        log(f"捷徑解析失敗：{f['name']}: {e}", "錯誤")
                else:
                    log(f"捷徑無 targetId，略過：{f['name']}", "Google Drive")
            else:
                resolved_batch.append(f)

        all_files.extend(resolved_batch)
        log(f"第 {page_num} 頁取得 {len(batch)} 個檔案，累計 {len(all_files)} 個", "Google Drive")

        page_token = response.get("nextPageToken")
        if not page_token:
            break

    log(f"分頁完成，共取得 {len(all_files)} 個檔案", "Google Drive")
    return all_files


def download_files_from_gdrive(folder_id, download_path):
    log(f"開始從 Google Drive 下載文件到: {download_path}", "Google Drive")

    if os.path.exists(download_path):
        shutil.rmtree(download_path)
    os.makedirs(download_path, exist_ok=True)

    files = list_all_files_from_gdrive(folder_id)
    log(f"共找到 {len(files)} 個檔案，開始下載...", "Google Drive")

    DIRECT_DOWNLOAD_TYPES = {
        'application/pdf',
        'text/plain',
        'text/markdown',
        'application/json',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'application/msword',
    }

    EXPORT_TYPES = {
        'application/vnd.google-apps.document': (
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document', '.docx'
        ),
        'application/vnd.google-apps.spreadsheet': (
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', '.xlsx'
        ),
    }

    MAX_RETRIES = 3

    for idx, file in enumerate(files, 1):
        file_id   = file['id']
        mime_type = file['mimeType']
        file_name = file['name'].replace("/", "-").replace("\\", "-")

        if mime_type in EXPORT_TYPES:
            export_mime, ext = EXPORT_TYPES[mime_type]
            if not file_name.lower().endswith(ext):
                file_name += ext
            file_path = os.path.join(download_path, file_name)
            log(f"[{idx}/{len(files)}] 匯出 Google Workspace 格式 -> {file_name}", "Google Drive")

            for attempt in range(1, MAX_RETRIES + 1):
                try:
                    request = drive_service.files().export_media(fileId=file_id, mimeType=export_mime)
                    with io.FileIO(file_path, 'wb') as fh:
                        downloader = MediaIoBaseDownload(fh, request)
                        done = False
                        while not done:
                            status, done = downloader.next_chunk()
                    log(f"    OK 匯出完成: {file_name}", "Google Drive")
                    break
                except Exception as e:
                    if attempt < MAX_RETRIES:
                        log(f"    RETRY [{attempt}/{MAX_RETRIES}] 匯出失敗，稍後重試 {file_name}: {e}", "Google Drive")
                        import time; time.sleep(3 * attempt)
                    else:
                        log(f"    FAIL 匯出失敗（已重試 {MAX_RETRIES} 次）{file_name}: {e}", "錯誤")

        elif mime_type in DIRECT_DOWNLOAD_TYPES:
            file_path = os.path.join(download_path, file_name)
            log(f"[{idx}/{len(files)}] 正在下載: {file_name}", "Google Drive")

            for attempt in range(1, MAX_RETRIES + 1):
                try:
                    request = drive_service.files().get_media(fileId=file_id)
                    with io.FileIO(file_path, 'wb') as fh:
                        downloader = MediaIoBaseDownload(fh, request)
                        done = False
                        while not done:
                            status, done = downloader.next_chunk()
                    log(f"    OK 下載完成: {file_name}", "Google Drive")
                    break
                except Exception as e:
                    if attempt < MAX_RETRIES:
                        log(f"    RETRY [{attempt}/{MAX_RETRIES}] 下載失敗，稍後重試 {file_name}: {e}", "Google Drive")
                        import time; time.sleep(3 * attempt)
                    else:
                        log(f"    FAIL 下載失敗（已重試 {MAX_RETRIES} 次）{file_name}: {e}", "錯誤")
        else:
            log(f"[{idx}/{len(files)}] 跳過不支援格式: {file_name} ({mime_type})", "Google Drive")

    log("所有文件下載完成", "Google Drive")
    return download_path

# ========== FastPDFLoader ==========

from langchain_core.documents import Document
import fitz  # PyMuPDF

class FastPDFLoader:
    def __init__(self, file_path):
        self.file_path = file_path

    def load(self):
        log(f"正在解析 PDF: {os.path.basename(self.file_path)}", "PDF解析")
        doc = fitz.open(self.file_path)
        text = "\n".join([page.get_text() for page in doc])
        log(f"PDF 解析完成，共 {len(doc)} 頁", "PDF解析")
        return [Document(page_content=text, metadata={"source": self.file_path})]

# ========== JSONLoader ==========

class JSONLoader:
    def __init__(self, file_path: str):
        self.file_path = file_path

    def _flatten(self, obj, prefix="") -> str:
        lines = []
        if isinstance(obj, dict):
            for k, v in obj.items():
                key_str = f"{prefix}{k}" if prefix else str(k)
                lines.append(self._flatten(v, prefix=f"{key_str}."))
        elif isinstance(obj, list):
            for i, item in enumerate(obj):
                lines.append(self._flatten(item, prefix=f"{prefix}[{i}]."))
        else:
            key_label = prefix.rstrip(".")
            lines.append(f"{key_label}: {obj}")
        return "\n".join(lines)

    def load(self) -> list:
        log(f"正在解析 JSON: {os.path.basename(self.file_path)}", "JSON解析")
        docs = []
        try:
            with open(self.file_path, "r", encoding="utf-8") as f:
                data = json.load(f)

            if isinstance(data, list):
                records = data
            elif isinstance(data, dict):
                wrapped = None
                for wrap_key in ("data", "records", "items", "results", "list"):
                    if wrap_key in data and isinstance(data[wrap_key], list):
                        wrapped = data[wrap_key]
                        break
                records = wrapped if wrapped is not None else [data]
            else:
                records = [{"value": data}]

            for idx, record in enumerate(records):
                text = self._flatten(record)
                if text.strip():
                    docs.append(Document(
                        page_content=text,
                        metadata={"source": self.file_path, "record_index": idx}
                    ))
            log(f"JSON 解析完成，共產生 {len(docs)} 個 Document", "JSON解析")
        except Exception as e:
            log(f"JSON 解析失敗: {e}", "錯誤")
        return docs

# ========== 全域向量資料庫快取 ==========

_shared_vectorstore = None
_shared_retriever   = None

# ========== 建立共用向量資料庫 ==========

def create_shared_vectorstore(force_rebuild: bool = False):
    global _shared_vectorstore

    if not force_rebuild and _shared_vectorstore is not None:
        log("使用記憶體快取的共用向量資料庫", "向量資料庫")
        return _shared_vectorstore

    vectorstore_path = f"breeze_vectorstores/{SHARED_VECTORSTORE_ID}"

    if force_rebuild and os.path.exists(vectorstore_path):
        shutil.rmtree(vectorstore_path)
        _shared_vectorstore = None

    if os.path.exists(vectorstore_path):
        log("發現現有共用向量資料庫，正在從磁碟載入...", "向量資料庫")
        embedding_model = HuggingFaceEmbeddings(model_name="BAAI/bge-m3")
        _shared_vectorstore = FAISS.load_local(
            vectorstore_path, embedding_model, allow_dangerous_deserialization=True
        )
        log("共用向量資料庫從磁碟載入完成", "向量資料庫")
        return _shared_vectorstore

    log("=== 開始建立新的共用向量資料庫 ===", "向量資料庫")
    download_path = f"drive_documents/{SHARED_VECTORSTORE_ID}"
    download_files_from_gdrive(GOOGLE_DRIVE_FOLDER_ID, download_path)

    md_loader = DirectoryLoader(
        download_path, glob="*.md",
        loader_cls=TextLoader,
        loader_kwargs={"autodetect_encoding": True}
    )
    md_docs = md_loader.load()
    log(f"載入 {len(md_docs)} 個 Markdown 文件", "文檔載入")

    pdf_docs = []
    for file in Path(download_path).glob("*.pdf"):
        pdf_docs.extend(FastPDFLoader(str(file)).load())
    log(f"載入 {len(pdf_docs)} 個 PDF 文件", "文檔載入")

    json_docs = []
    for file in Path(download_path).glob("*.json"):
        json_docs.extend(JSONLoader(str(file)).load())
    log(f"載入 {len(json_docs)} 個 JSON Document", "文檔載入")

    from docx import Document as DocxDocument
    word_docs = []
    for file in Path(download_path).glob("*.docx"):
        try:
            docx_file = DocxDocument(str(file))
            text = "\n".join([p.text for p in docx_file.paragraphs if p.text.strip()])
            if text.strip():
                word_docs.append(Document(page_content=text, metadata={"source": str(file)}))
        except Exception as e:
            log(f"Word 載入失敗 {file.name}: {e}", "錯誤")
    log(f"載入 {len(word_docs)} 個 Word 文件", "文檔載入")

    all_docs = md_docs + pdf_docs + json_docs + word_docs
    if not all_docs:
        log("沒有找到任何文件", "錯誤")
        return None

    # ★ 方向二：chunk_size 從 1000 縮小至 500，overlap 從 200 調整至 150
    #   避免單一 chunk 混入多個不相關服務項目，提升擷取精確度
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=150)
    all_splits = []
    with ThreadPoolExecutor(max_workers=8) as executor:
        for chunks in executor.map(lambda d: text_splitter.split_documents([d]), all_docs):
            all_splits.extend(chunks)
    log(f"文檔切割完成：共 {len(all_splits)} 個片段", "文檔切割")

    embedding_model = HuggingFaceEmbeddings(model_name="BAAI/bge-m3")
    vectorstore = FAISS.from_documents(all_splits, embedding_model)
    vectorstore.save_local(vectorstore_path)

    _shared_vectorstore = vectorstore
    log("=== 共用向量資料庫建立完成 ===", "向量資料庫")
    return _shared_vectorstore

# ========== 取得共用向量資料庫物件 ==========

def get_shared_vectorstore(force_rebuild: bool = False):
    """直接回傳向量資料庫物件，供自訂檢索邏輯（偽查詢擴充）使用"""
    return create_shared_vectorstore(force_rebuild=force_rebuild)

# ========== 取得共用 Retriever（供初始化與重置使用）==========

def get_shared_retriever(force_rebuild: bool = False):
    global _shared_retriever

    if not force_rebuild and _shared_retriever is not None:
        return _shared_retriever

    vectorstore = create_shared_vectorstore(force_rebuild=force_rebuild)
    if vectorstore is None:
        return None

    # ★ 方向二 / 六：lambda_mult 從 0.5 提高至 0.75，更偏重相關性
    _shared_retriever = vectorstore.as_retriever(
        search_type="mmr",
        search_kwargs={'k': 3, 'fetch_k': 20, 'lambda_mult': 0.75}
    )
    log("Retriever 建立完成 (k=3, fetch_k=20, lambda_mult=0.75)", "檢索器")
    return _shared_retriever

# ========== Whisper 語音辨識 ==========

log("正在載入 Whisper 語音辨識模型...", "系統")
audio_model = whisper.load_model("turbo")
log("Whisper 模型載入完成", "系統")

def trans_aac_to_text(aac_path):
    try:
        result = audio_model.transcribe(aac_path, language="zh")
        return result["text"]
    except Exception as e:
        log(f"語音辨識失敗: {e}", "錯誤")
        return "Could not understand audio"

cc = OpenCC('s2twp')  # 簡體轉台灣正體

# ========== 語音偏好設定 ==========

user_voice_pref: dict[str, str] = {}  # user_id → "female" | "male"

VOICE_MAP = {
    "female": "zh-TW-HsiaoChenNeural",   # 小晨（女）
    "male":   "zh-TW-YunJheNeural",      # 雲哲（男）
}

def get_user_voice(user_id: str) -> str:
    pref = user_voice_pref.get(user_id, "female")
    return VOICE_MAP[pref]

async def synthesize_speech(text: str, voice: str, mp3_path: str):
    """使用 edge-tts 合成語音並儲存為 mp3"""
    communicate = edge_tts.Communicate(text, voice)
    await communicate.save(mp3_path)

# ========== Chain 建立函數 ==========

def create_extract_chain():
    """從 RAG 撈回的 chunk 片段中，根據問題摘要出相關資訊"""
    llama_model = ChatOllama(model="jcai/breeze-7b-instruct-v1_0:q4_0", temperature=0)

    extract_prompt = ChatPromptTemplate.from_messages([
        ("human",
         """你是一個只根據文件內容回答問題的助理。

嚴格規則：
1. 只能根據下方「文件內容」回答，不可使用自己的知識。
2. 禁止使用簡體字，必須使用繁體中文，禁止中英文夾雜。
3. 若文件中沒有與問題相關的資訊，請只回答：「資料庫中查無相關資料」，不可補充任何其他內容。
4. 不可猜測或自行補充任何資訊。
5. 只摘要與問題直接相關的內容，與問題無關的段落一律略過，不可因補充背景而帶入不相關內容。
6. 嚴格禁止做出任何疾病診斷，不可將輕微症狀與嚴重疾病直接關聯，不可誇大症狀嚴重性，只能提供一般性資訊。

文件內容：
{context}

問題：
{question}

摘要回答：
""")
    ])

    return extract_prompt | llama_model | StrOutputParser()


def create_combine_chain():
    """將 RAG 摘要潤飾為口語化、有溫度的回覆"""
    llama_model = ChatOllama(model="jcai/breeze-7b-instruct-v1_0:q4_0", temperature=1)

    combine_prompt = ChatPromptTemplate.from_template(
        """你是「健康小幫手」，一位在社工機構服務的親切健康諮詢志工。
說話風格溫柔、口語化，像朋友在聊天，說話有溫度且充滿關心。你的名字叫小伊，若有需要請以小伊自稱。

嚴格規則：
1. 只能根據「資料庫參考回答」的內容進行回覆，不可自行補充或創造任何知識。
2. 必須使用繁體中文，禁止使用簡體字，禁止中英文夾雜。
3. 字數控制在 250 字以內。
4. 若有對話歷史，請自然地銜接上下文，讓對話有連貫感。
5. 請適當斷句與換行，確保語意清晰、易於閱讀。
6. 可依語境適度加入表情符號（如 😉、✅、😄、⚠️、👍），每則最多 1～3 個，避免過度使用。
7. 開頭關懷句請從不同語氣中選擇，不要每次都使用相同句型。
8. 結尾關懷句請自然變化，不要與前次回答重複。
9. 嚴格禁止做出任何疾病診斷或暗示，不可將輕微症狀誇大為嚴重疾病，只能提供一般性建議並鼓勵就醫。

回覆結構（依序）：
a. 【開場關懷】視情況加入一句貼近問題的關心語
b. 【核心回答】根據參考回答清楚說明，語氣口語自然
c. 【結尾關懷】加入一句鼓勵或關懷語

參考範例：
範例 1：
問題：我最近都睡不好怎麼辦
資料庫參考回答：建議規律作息、睡前避免使用電子產品，必要時就醫
回覆：
睡不好真的會影響整天的精神呢。
建議您可以試著維持固定的作息時間，睡前避免滑手機，如果持續沒有改善，也可以考慮尋求醫師協助喔。
希望您能慢慢找回好的睡眠品質！👍

範例 2：
問題：我最近壓力很大怎麼辦
資料庫參考回答：建議適度休息、運動或與他人傾訴
回覆：
感覺您最近背負了不少壓力，先給自己一點空間也很重要。
建議可以適度安排休息時間，做些放鬆的活動，像是運動或和信任的人聊聊。
今天也辛苦您了，別忘了給自己一點放鬆的時間！😄

---
對話歷史（最近幾輪）：
{history}

目前問題：
{question}

資料庫參考回答：
{rag_answer}

自然回覆："""
    )
    return combine_prompt | llama_model | StrOutputParser()


def create_fallback_chain():
    """RAG 查無資料時，改由模型自身知識生成回答（保留小伊人設）
    ★ 方向一：移除「📋 知識庫未收錄」免責開頭，改由 process_message 依意圖動態決定是否加入
    """
    llama_model = ChatOllama(model="jcai/breeze-7b-instruct-v1_0:q4_0", temperature=1)

    fallback_prompt = ChatPromptTemplate.from_template(
        """你是「健康小幫手」，一位在社工機構服務的親切健康諮詢志工。
說話風格溫柔、口語化，像朋友在聊天，說話有溫度且充滿關心。你的名字叫小伊，若有需要請以小伊自稱。

規則：
1. 根據你的醫療健康常識或心理支持知識回答問題，語氣口語自然。
2. 必須使用繁體中文，禁止使用簡體字，禁止中英文夾雜。
3. 字數控制在 250 字以內。
4. 若有對話歷史，請自然地銜接上下文，讓對話有連貫感。
5. 請適當斷句與換行，確保語意清晰、易於閱讀。
6. 可依語境適度加入表情符號，每則最多 1～3 個。
7. 回覆結尾請提醒：「如有需要，建議您諮詢專業醫師或相關機構，以獲得更準確的建議 🙏」

回覆結構（依序）：
a. 【開場關懷】一句貼近問題的關心語
b. 【核心回答】根據一般常識清楚說明
c. 【結尾提醒】固定結尾（規則 7）

---
對話歷史（最近幾輪）：
{history}

目前問題：
{question}

自然回覆："""
    )
    return fallback_prompt | llama_model | StrOutputParser()


# ★ 方向二 / 四：合併分類 + 代名詞替換為一次 Ollama 呼叫，減少推論次數
def create_classify_rewrite_chain():
    """合併問句分類與代名詞替換，單次 Ollama 呼叫同時完成兩個任務"""
    llama_model = ChatOllama(model="jcai/breeze-7b-instruct-v1_0:q4_0", temperature=0)

    prompt = ChatPromptTemplate.from_template(
        """你是問句分析助理，請根據對話歷史完成以下兩個任務。

【任務一：分類】
判斷目前問題是「追問」還是「提問」：
- 追問：含「這裡/那裡/它/他/她/這個/那個/這件事/這樣/那樣」等代名詞，需參考前文才能理解
- 提問：問句本身完整，不需前文（縮寫如「北市復康」「健保」「長照」不算代名詞）

【任務二：改寫】
- 若為追問：將代名詞替換為對話歷史中對應的具體名詞，其餘文字一律不動，不可新增原文沒有的資訊
- 若為提問：直接輸出原句，不做任何修改

輸出格式（只輸出這兩行，不要其他文字）：
類型：追問 或 提問
問句：[改寫後的問句 或 原句]

範例：
對話歷史：使用者：北市復康的費用？ 助理：費用為計程車費率的三分之一
目前問題：那他的服務時間呢？
類型：追問
問句：北市復康的服務時間呢？

對話歷史：（無）
目前問題：長照補助怎麼申請？
類型：提問
問句：長照補助怎麼申請？

---
對話歷史：
{history}

目前問題：{question}

輸出："""
    )
    return prompt | llama_model | StrOutputParser()


# ★ 方向二：偽查詢生成 chain，提升 RAG 檢索精確度
def create_pseudo_query_chain():
    """根據使用者問題生成假設性答案片段，用以擴充 FAISS 檢索範圍"""
    llama_model = ChatOllama(model="jcai/breeze-7b-instruct-v1_0:q4_0", temperature=0)

    prompt = ChatPromptTemplate.from_template(
        """根據以下問題，用繁體中文生成 2 個可能出現在相關文件中的簡短答案片段（每個不超過 30 字）。
只輸出片段，每行一個，不要編號或任何說明文字。

問題：{question}

片段："""
    )
    return prompt | llama_model | StrOutputParser()


# ========== 模型池（Queue Pool）+ Ollama 限流 Semaphore ==========
#
# ★ 方向四：
#   - classify_chain + rewrite_chain 合併為 classify_rewrite_chain（減少一次 Ollama 呼叫）
#   - 新增 pseudo_query_chain_pool（偽查詢生成）
#   - _ollama_semaphore 維持 1，確保 CPU 推論穩定
#
# Pool size 建議：
#   - CPU 推論（q4_0）：POOL_SIZE=2, Semaphore=1
#   - 記憶體不足時：POOL_SIZE=1, Semaphore=1
# =========================================================

POOL_SIZE = 2

extract_chain_pool        = asyncio.Queue()
combine_chain_pool        = asyncio.Queue()
fallback_chain_pool       = asyncio.Queue()
classify_rewrite_chain_pool = asyncio.Queue()  # ★ 合併後的分類+改寫池
pseudo_query_chain_pool   = asyncio.Queue()    # ★ 偽查詢生成池

# Ollama 同時只允許 1 個推論，其餘排隊等待
_ollama_semaphore = asyncio.Semaphore(1)


async def init_pools():
    """啟動時預先建立所有 chain 實例並放入各自的 Pool"""
    log(f"正在初始化模型池（每個 chain x{POOL_SIZE} 個實例）...", "模型池")
    for i in range(POOL_SIZE):
        await extract_chain_pool.put(create_extract_chain())
        await combine_chain_pool.put(create_combine_chain())
        await fallback_chain_pool.put(create_fallback_chain())
        await classify_rewrite_chain_pool.put(create_classify_rewrite_chain())
        await pseudo_query_chain_pool.put(create_pseudo_query_chain())
        log(f"  實例 {i+1}/{POOL_SIZE} 建立完成", "模型池")
    log("模型池初始化完成", "模型池")


async def async_invoke_from_pool(pool: asyncio.Queue, input_data: dict) -> str:
    """
    從 Pool 取出 chain → 用 Semaphore 限流 → 執行推論 → 放回 Pool
    finally 保證 chain 一定放回，即使推論拋出例外也不會縮水
    """
    chain = await pool.get()
    try:
        async with _ollama_semaphore:
            result = await asyncio.to_thread(chain.invoke, input_data)
            if result is None:
                raise RuntimeError("chain.invoke 回傳 None")
    finally:
        await pool.put(chain)
    return result


# ========== 問句分類與改寫（合併版）==========

async def classify_and_rewrite(user_question: str, history: list) -> str:
    """
    ★ 方向二 / 四：合併分類與代名詞替換為單次 Ollama 呼叫
    - 解析輸出的「類型：」與「問句：」兩行
    - 安全保護：改寫後句子長度若超過原句 50%，退回原句（防止模型過度發揮）
    - log 輸出改寫前後對比，方便日後驗證
    """
    if not history:
        log("無對話歷史，直接視為提問", "問句分析")
        return user_question

    history_text = format_history_text(history[-2:])

    raw_output = await async_invoke_from_pool(classify_rewrite_chain_pool, {
        "history": history_text,
        "question": user_question,
    })

    # 解析輸出
    classify_result = ""
    rewritten = user_question
    for line in raw_output.strip().split('\n'):
        line = line.strip()
        if line.startswith("類型："):
            classify_result = line.replace("類型：", "").strip()
        elif line.startswith("問句："):
            candidate = line.replace("問句：", "").strip()
            if candidate:
                rewritten = candidate

    log(f"問句分類：{classify_result}｜{user_question} → {rewritten}", "問句分析")

    # ★ 安全保護：改寫句超過原句長度 50% 時退回原句
    if len(rewritten) > len(user_question) * 1.5:
        log(f"改寫句過長（{len(rewritten)} vs {len(user_question)}），退回原句", "問句分析")
        return user_question

    return rewritten


# ========== 偽查詢擴充檢索 ==========

async def retrieve_with_pseudo_queries(vectorstore, standalone_question: str) -> list:
    """
    ★ 方向二：Pseudo Query Reranker
    1. 用 Breeze 生成 2 個假設性答案片段
    2. 每個片段對 FAISS 做 similarity_search（快速，非 Ollama）
    3. 與原始 MMR 檢索結果合併去重
    4. 最終回傳最多 5 個相關 chunk
    """
    # Step 1：生成偽查詢片段（Ollama 呼叫）
    try:
        pseudo_text = await async_invoke_from_pool(pseudo_query_chain_pool, {
            "question": standalone_question
        })
        pseudo_queries = [q.strip() for q in pseudo_text.strip().split('\n') if q.strip()][:2]
        log(f"偽查詢生成：{pseudo_queries}", "偽查詢")
    except Exception as e:
        log(f"偽查詢生成失敗，使用原始問句: {e}", "偽查詢")
        pseudo_queries = []

    # Step 2：原始 MMR 檢索（使用 vectorstore 直接建立 retriever）
    retriever = vectorstore.as_retriever(
        search_type="mmr",
        search_kwargs={'k': 3, 'fetch_k': 20, 'lambda_mult': 0.75}
    )
    main_docs = await asyncio.to_thread(retriever.invoke, standalone_question)

    # Step 3：偽查詢 FAISS 擴充檢索（不走 Ollama，純向量計算）
    all_docs = list(main_docs)
    seen_content = {doc.page_content[:100] for doc in all_docs}

    for pq in pseudo_queries:
        try:
            pq_docs = await asyncio.to_thread(vectorstore.similarity_search, pq, k=2)
            for doc in pq_docs:
                key = doc.page_content[:100]
                if key not in seen_content:
                    seen_content.add(key)
                    all_docs.append(doc)
        except Exception as e:
            log(f"偽查詢擴充失敗: {e}", "偽查詢")

    log(f"擴充檢索完成，共 {len(all_docs)} 個相關片段（上限 5 個）", "偽查詢")
    return all_docs[:5]


# ========== 主要處理流程 ==========

async def full_pipeline(user_id: str, user_question: str, standalone_question: str):
    log("========== 開始 AI 問答流程 ==========", "AI問答")

    vectorstore = get_shared_vectorstore()
    if vectorstore is None:
        return {
            "rag_answer": "無法載入知識庫",
            "final_response": "抱歉，小伊目前無法載入知識庫，請稍後再試。",
            "sources": [],
            "is_fallback": False,
        }

    history = get_user_history(user_id)
    history_text = format_history_text(history) if history else "（無對話歷史）"

    # ── Step 1：偽查詢擴充檢索 ──
    log(f"RAG 檢索問句：{standalone_question}", "AI問答")
    source_docs = await retrieve_with_pseudo_queries(vectorstore, standalone_question)

    # ── Step 2：直接使用 chunk 片段作為 context（方向四：不再載入全文）──
    SEP = "=" * 60
    DIV = "-" * 60

    sources = []
    chunk_context_parts = []
    seen_chunks = set()

    for doc in source_docs:
        fname = os.path.basename(doc.metadata.get('source', ''))
        content = doc.page_content.strip()
        key = content[:100]

        if key not in seen_chunks:
            seen_chunks.add(key)
            chunk_context_parts.append(f"【{fname}】\n{content}")

        snippet = content.replace("\n", " ")[:120]
        sources.append({"file": fname, "snippet": snippet})

    print(f"\n{SEP}")
    print(f"  [RAG 檢索結果]")
    print(DIV)
    q_display = standalone_question[:100] + ("..." if len(standalone_question) > 100 else "")
    print(f"  問題輸入（獨立問句）：{q_display}")
    print(DIV)
    print(f"  參考來源 chunk（共 {len(chunk_context_parts)} 個）：")
    for i, s in enumerate(sources[:5], 1):
        print(f"  [{i}] {s['file']}")
        print(f"       -> {s['snippet']}...")

    if not chunk_context_parts:
        print(f"       （無相關文件片段）")
        print(f"{SEP}\n")
        log("無相關文件片段，改用 Fallback 模型自身知識生成回答...", "AI問答")
        fallback_response = await async_invoke_from_pool(fallback_chain_pool, {
            "history": history_text,
            "question": user_question,
        })
        fallback_response = cc.convert(fallback_response)
        return {
            "rag_answer": "無相關文件片段",
            "final_response": fallback_response,
            "sources": sources,
            "is_fallback": True,
        }

    # ── Step 3：extract_chain 從 chunk 摘要相關內容 ──
    full_context = ("\n\n" + "=" * 40 + "\n\n").join(chunk_context_parts)
    log("正在從 chunk 中摘要相關內容（排隊等待模型池）...", "AI問答")

    rag_answer = await async_invoke_from_pool(extract_chain_pool, {
        "context": full_context,
        "question": standalone_question,
    })

    print(DIV)
    print(f"  RAG 摘要回答：")
    for line in rag_answer.strip().split("\n"):
        print(f"    {line}")
    print(f"{SEP}\n")

    # ── 查無資料 → fallback chain ──
    if "資料庫中查無相關資料" in rag_answer:
        log("RAG 查無相關資料，改用 Fallback 模型自身知識生成回答...", "AI問答")

        fallback_response = await async_invoke_from_pool(fallback_chain_pool, {
            "history": history_text,
            "question": user_question,
        })
        fallback_response = cc.convert(fallback_response)

        print(f"\n{SEP}")
        print(f"  [Fallback 回覆（模型自身知識）]")
        print(DIV)
        for line in fallback_response.strip().split("\n"):
            print(f"  {line}")
        print(f"{SEP}\n")

        return {
            "rag_answer": rag_answer,
            "final_response": fallback_response,
            "sources": sources,
            "is_fallback": True,
        }

    # ── Step 4：combine_chain 潤飾為口語化回覆 ──
    log("正在整合最終回覆（排隊等待模型池）...", "AI問答")

    final_response = await async_invoke_from_pool(combine_chain_pool, {
        "history": history_text,
        "question": user_question,
        "rag_answer": rag_answer,
    })
    final_response = cc.convert(final_response)

    print(f"\n{SEP}")
    print(f"  [最終回覆]")
    print(DIV)
    for line in final_response.strip().split("\n"):
        print(f"  {line}")
    print(f"{SEP}\n")

    log("最終回覆生成完成", "AI問答")
    return {
        "rag_answer": rag_answer,
        "final_response": final_response,
        "sources": sources,
        "is_fallback": False,
    }


# ========== FastAPI 啟動事件 ==========

@app.on_event("startup")
async def startup_event():
    log("=" * 50, "系統")
    log("LINE Bot 小伊正在啟動...", "系統")
    log(f"ngrok URL: {ngrok_url}", "系統")

    log("【啟動預載】正在初始化 Google Sheets...", "系統")
    await asyncio.to_thread(init_sheets)

    log("【啟動預載】開始初始化向量資料庫...", "系統")
    try:
        await asyncio.to_thread(get_shared_retriever)
        log("【啟動預載】向量資料庫初始化完成", "系統")
    except Exception as e:
        log(f"【啟動預載】初始化失敗: {e}", "錯誤")

    log("【啟動預載】開始初始化模型池...", "系統")
    try:
        await init_pools()
        log("【啟動預載】模型池初始化完成，系統已就緒！", "系統")
    except Exception as e:
        log(f"【啟動預載】模型池初始化失敗: {e}", "錯誤")

    log("=" * 50, "系統")
    log("LINE Bot 小伊已啟動完成！", "系統")
    log("=" * 50, "系統")


# 儲存主事件迴圈
main_loop = asyncio.get_event_loop()

@app.post("/callback")
async def callback(request: Request):
    log("收到 LINE Webhook 請求", "Webhook")
    signature = request.headers.get("X-Line-Signature", "")
    body = await request.body()
    body_text = body.decode("utf-8")

    try:
        await asyncio.to_thread(handler.handle, body_text, signature)
    except Exception as e:
        log(f"Webhook 處理錯誤: {e}", "錯誤")
        raise HTTPException(status_code=500, detail="Handler error")

    return "OK"

# ========== 文字 / 語音訊息統一處理 ==========

@handler.add(MessageEvent)
def unified_handler(event):
    message = event.message
    user_id = event.source.user_id

    if isinstance(message, TextMessageContent):
        user_text = message.text
        log(f"訊息類型: 文字 | 內容: {user_text}", "訊息接收")
        try:
            asyncio.run_coroutine_threadsafe(
                process_message(event, user_id, user_text), main_loop
            )
        except Exception as e:
            log(f"訊息排程失敗: {e}", "錯誤")

    elif isinstance(message, AudioMessageContent):
        log("訊息類型: 語音", "訊息接收")
        with ApiClient(configuration) as api_client:
            blob_api = MessagingApiBlob(api_client)
            content = blob_api.get_message_content(message_id=message.id)
            aac_path = f'./audio/{user_id}.aac'
            with open(aac_path, 'wb') as f:
                f.write(content)

        user_text = trans_aac_to_text(aac_path)
        try:
            asyncio.run_coroutine_threadsafe(
                process_message(event, user_id, user_text), main_loop
            )
        except Exception as e:
            log(f"訊息排程失敗: {e}", "錯誤")

# ========== Quick Reply ==========

def build_end_quick_reply(user_id: str = "") -> QuickReply:
    current = user_voice_pref.get(user_id, "female")
    switch_label = "🔊 切換男聲" if current == "female" else "🔊 切換女聲"
    switch_text  = "！切換男聲！" if current == "female" else "！切換女聲！"
    return QuickReply(items=[
        QuickReplyItem(action=MessageAction(label=switch_label, text=switch_text)),
        QuickReplyItem(action=MessageAction(label="👋 結束使用", text="！結束使用！")),
    ])

# ========== Flex Message：五星評分 ==========

def build_rating_flex_message() -> FlexMessage:
    flex_body = {
        "type": "bubble",
        "body": {
            "type": "box",
            "layout": "vertical",
            "spacing": "sm",
            "contents": [
                {"type": "text", "text": "感謝您的使用！", "weight": "bold", "size": "lg", "align": "center"},
                {"type": "text", "text": "可以幫小伊評分一下這次的服務嗎〜😉", "size": "sm", "color": "#888888", "align": "center", "margin": "sm"},
                {"type": "separator", "margin": "md"},
                *[
                    {
                        "type": "box", "layout": "vertical",
                        "margin": "md" if i == 0 else "sm",
                        "backgroundColor": color,
                        "cornerRadius": "4px", "paddingAll": "12px",
                        "action": {"type": "message", "label": label, "text": f"評分：{i+1} 星"},
                        "contents": [{"type": "text", "text": label, "align": "center", "color": "#000000", "weight": "bold"}]
                    }
                    for i, (label, color) in enumerate([
                        ("★☆☆☆☆ 1 星", "#FFD700"),
                        ("★★☆☆☆ 2 星", "#FFD700"),
                        ("★★★☆☆ 3 星", "#FFD700"),
                        ("★★★★☆ 4 星", "#FFD700"),
                        ("★★★★★ 5 星", "#FFD700"),
                    ])
                ]
            ]
        }
    }
    return FlexMessage(
        alt_text="請為小伊本次的服務評分⭐",
        contents=FlexContainer.from_dict(flex_body)
    )

# ========== 語音合成（reply 版）==========

async def generate_audio_message(event, reply: str, user_id: str):
    """使用 reply_token 回覆，適合即時、快速的回應（如等待中、閒聊）"""
    voice    = get_user_voice(user_id)
    mp3_path = f"./audio/{user_id}_reply.mp3"
    # ★ 方向三：TTS 傳入過濾 emoji 後的文字，TextMessage 仍保留原始 emoji
    await synthesize_speech(strip_emoji(reply), voice, mp3_path)

    reply_m4a_path = f"./audio/{user_id}_reply.m4a"
    sound = AudioSegment.from_file(mp3_path)
    sound.export(reply_m4a_path, format="mp4", codec="aac")

    audio_info   = mediainfo(mp3_path)
    duration_ms  = int(float(audio_info['duration']) * 1000)
    audio_url    = f"{ngrok_url}/audio/{user_id}_reply.m4a"

    with ApiClient(configuration) as api_client:
        line_bot_api = MessagingApi(api_client)
        line_bot_api.reply_message(
            ReplyMessageRequest(
                reply_token=event.reply_token,
                messages=[
                    TextMessage(text=reply),
                    AudioMessage(
                        original_content_url=audio_url,
                        duration=duration_ms,
                        quick_reply=build_end_quick_reply(user_id)
                    )
                ]
            )
        )
    log("Reply 回覆已發送（文字 + 語音 + Quick Reply）", "訊息發送")

# ========== 語音合成（push 版）==========

async def push_audio_message(user_id: str, reply: str):
    """使用 push_message 傳送，不依賴 reply_token，適合耗時的 AI 回覆"""
    try:
        voice    = get_user_voice(user_id)
        mp3_path = f"./audio/{user_id}_reply.mp3"
        # ★ 方向三：TTS 傳入過濾 emoji 後的文字，TextMessage 仍保留原始 emoji
        await synthesize_speech(strip_emoji(reply), voice, mp3_path)

        reply_m4a_path = f"./audio/{user_id}_reply.m4a"
        sound = AudioSegment.from_file(mp3_path)
        sound.export(reply_m4a_path, format="mp4", codec="aac")

        audio_info  = mediainfo(mp3_path)
        duration_ms = int(float(audio_info['duration']) * 1000)
        audio_url   = f"{ngrok_url}/audio/{user_id}_reply.m4a"

        with ApiClient(configuration) as api_client:
            line_bot_api = MessagingApi(api_client)
            await asyncio.to_thread(
                line_bot_api.push_message,
                PushMessageRequest(
                    to=user_id,
                    messages=[
                        TextMessage(text=reply),
                        AudioMessage(
                            original_content_url=audio_url,
                            duration=duration_ms,
                            quick_reply=build_end_quick_reply(user_id)
                        )
                    ]
                )
            )
        log("Push 回覆已發送（文字 + 語音 + Quick Reply）", "訊息發送")
    except Exception as e:
        log(f"push_audio_message 失敗: {e}", "錯誤")

# ========== 主訊息處理流程 ==========

async def process_message(event, user_id: str, msg: str):
    log(f"\n{'=' * 60}", "訊息處理")
    log(f"開始處理使用者訊息: {msg}", "訊息處理")

    try:
        # ── 特殊指令：重新整理知識庫 ──
        if msg.strip() == "！重置小伊的知識庫！":
            await asyncio.to_thread(get_shared_retriever, True)
            await generate_audio_message(event, "小伊的知識庫已完成更新✅", user_id)
            return

        # ── 結束使用 ──
        if msg.strip() == "！結束使用！":
            with ApiClient(configuration) as api_client:
                MessagingApi(api_client).reply_message(
                    ReplyMessageRequest(
                        reply_token=event.reply_token,
                        messages=[build_rating_flex_message()]
                    )
                )
            return

        # ── 切換語音偏好 ──
        if msg.strip() in ("！切換男聲！", "！切換女聲！"):
            pref  = "male"   if "男" in msg else "female"
            label = "男聲👦🏻" if pref == "male" else "女聲👧🏻"
            user_voice_pref[user_id] = pref
            log(f"語音切換：{user_id[:8]} -> {label}", "語音設定")
            reply = f"已切換為{label}，之後的回覆都會用這個聲音喔 😊"
            await push_audio_message(user_id, reply)
            return

        # ── 評分回覆 ──
        if msg.strip().startswith("評分：") and msg.strip().endswith("星"):
            score  = msg.strip()
            last_q = user_last_question.get(user_id, "")
            await asyncio.to_thread(save_rating_to_sheet, user_id, score, last_q)

            with ApiClient(configuration) as api_client:
                MessagingApi(api_client).reply_message(
                    ReplyMessageRequest(
                        reply_token=event.reply_token,
                        messages=[TextMessage(
                            text=f"感謝您給了 {score}！您的回饋對小伊非常重要，期待下次再為您服務 😄"
                        )]
                    )
                )
            return

        # ── 意圖分類 ──
        history = get_user_history(user_id)
        intent  = classify_intent(msg, history)
        log(f"意圖分類: {intent}", "意圖分類")

        # ── 閒聊：直接回固定模板，不走 RAG ──
        if intent == "chitchat":
            reply = get_chitchat_response(msg)
            await generate_audio_message(event, reply, user_id)
            return

        # ── 一般問題 / 追問 / 情緒支持：走 RAG ──
        log("開始生成 AI 回答...", "訊息處理")

        # ★ 方向四：立即用 reply_token 回覆等待訊息，讓使用者知道系統有收到
        with ApiClient(configuration) as api_client:
            try:
                MessagingApi(api_client).reply_message(
                    ReplyMessageRequest(
                        reply_token=event.reply_token,
                        messages=[TextMessage(text="⏳ 小伊正在查詢中，請稍候...")]
                    )
                )
                log("等待訊息已發送", "訊息發送")
            except Exception as e:
                log(f"等待訊息發送失敗（不影響主流程）: {e}", "訊息發送")

        add_to_history(user_id, "user", msg)
        user_last_question[user_id] = msg

        # Step 1：問句改寫（合併版 chain）
        standalone = await classify_and_rewrite(msg, history)
        standalone_question = standalone + "，請用繁體中文回答"
        user_question       = msg        + "，請用繁體中文回答"

        # Step 2：RAG 流程
        result = await full_pipeline(user_id, user_question, standalone_question)

        reply = result["final_response"]

        # ★ 方向一：依意圖決定是否加入「知識庫未收錄」免責開頭
        #   - 情緒支持 + fallback → 不加，保持溫暖語氣
        #   - 知識問題 + fallback → 加免責說明
        if result.get("is_fallback"):
            if intent != "emotional_support":
                reply = (
                    "📋 小伊的知識庫目前沒有收錄這個問題的資料，"
                    "以下是小伊根據一般常識整理的參考資訊，僅供參考喔！\n\n"
                    + reply
                )
            log(f"Fallback 回覆，寫入查無資料記錄 | 意圖: {intent} | 問題: {msg}", "Fallback")
            await asyncio.to_thread(
                save_no_result_log,
                user_id,
                msg,
                result["rag_answer"],
                reply,
            )

        # 更新對話歷史並發送回覆
        add_to_history(user_id, "assistant", reply)
        await push_audio_message(user_id, reply)
        log("訊息處理完成", "訊息處理")

    except Exception as e:
        log(f"處理訊息時發生錯誤: {e}", "錯誤")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    log("正在啟動 Uvicorn 伺服器...", "系統")
    uvicorn.run("eden_linebot:app", host="0.0.0.0", port=8000, reload=False)
